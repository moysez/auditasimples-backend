import os
from datetime import datetime
from typing import Any, Dict
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _fmt_money(v: Any) -> str:
    try:
        n = float(v or 0.0)
    except Exception:
        n = 0.0
    return f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def _fmt_percent(frac: Any) -> str:
    try:
        n = float(frac or 0.0)
    except Exception:
        n = 0.0
    n = n * 100.0
    return f"{n:,.2f}%".replace(",", "X").replace(".", ",").replace("X", ".")

def _format_table_borders(table) -> None:
    tbl = table._tbl
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tbl.tblPr.append(tblBorders)

def _add_row(tbl, label: Any, value: Any) -> None:
    r = tbl.add_row().cells
    r[0].text = "" if label is None else str(label)
    r[1].text = "" if value is None else str(value)

def gerar_relatorio_fiscal(totals: Dict[str, Any], client_name: str = "Cliente", cnpj: str = "00.000.000/0000-00") -> str:
    totals = totals or {}
    tax = totals.get("tax_summary") or {}

    period_start = totals.get("period_start") or "---"
    period_end   = totals.get("period_end") or "---"

    documentos  = totals.get("documents", 0) or 0
    itens       = totals.get("items", 0) or 0
    valor_total = totals.get("total_value_sum", 0.0) or 0.0

    faturamento     = tax.get("faturamento", valor_total) or 0.0
    receita_excluida= tax.get("receita_excluida", 0.0) or 0.0
    base_corrigida  = tax.get("base_corrigida", faturamento - receita_excluida) or 0.0
    imposto_pago    = tax.get("imposto_pago", 0.0) or 0.0
    imposto_corrigido = tax.get("imposto_corrigido", 0.0) or 0.0
    economia        = tax.get("economia_estimada", 0.0) or 0.0
    aliquota_frac   = tax.get("aliquota_utilizada", 0.0) or 0.0

    produtos_duplicados = totals.get("produtos_duplicados") or []

    doc = Document()
    data_hoje = datetime.now().strftime("%d/%m/%Y")

    h = doc.add_heading("Relatório de Auditoria Fiscal — (Simples Nacional)", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"{client_name}  |  CNPJ: {cnpj}")
    p.runs[0].bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(f"Período auditado: {period_start} — {period_end}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3 = doc.add_paragraph(f"Data da análise: {data_hoje}")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # 1. Resumo Geral
    doc.add_heading("1. Resumo Geral", level=1)
    t_resumo = doc.add_table(rows=1, cols=2)
    t_resumo.style = "Table Grid"
    t_resumo.cell(0,0).text="Indicador"
    t_resumo.cell(0,1).text="Valor"
    for c in t_resumo.rows[0].cells: c.paragraphs[0].runs[0].bold = True
    _add_row(t_resumo, "Documentos analisados", documentos)
    _add_row(t_resumo, "Itens analisados", itens)
    _add_row(t_resumo, "Faturamento bruto (soma NF-e)", f"R$ {_fmt_money(faturamento)}")
    _format_table_borders(t_resumo)
    doc.add_paragraph("")

    # 2. Resumo Tributário
    doc.add_heading("2. Resumo Tributário", level=1)
    t_trib = doc.add_table(rows=1, cols=2)
    t_trib.style = "Table Grid"
    t_trib.cell(0,0).text="Descrição"
    t_trib.cell(0,1).text="Valor"
    for c in t_trib.rows[0].cells: c.paragraphs[0].runs[0].bold = True
    _add_row(t_trib, "Faturamento Bruto",            f"R$ {_fmt_money(faturamento)}")
    _add_row(t_trib, "Receita Monofásica Excluída", f"R$ {_fmt_money(receita_excluida)}")
    _add_row(t_trib, "Base Corrigida",              f"R$ {_fmt_money(base_corrigida)}")
    _add_row(t_trib, "Alíquota utilizada",           _fmt_percent(aliquota_frac))
    _add_row(t_trib, "Imposto Pago",                f"R$ {_fmt_money(imposto_pago)}")
    _add_row(t_trib, "Imposto Corrigido",           f"R$ {_fmt_money(imposto_corrigido)}")
    _add_row(t_trib, "Economia Estimada",           f"R$ {_fmt_money(economia)}")
    _format_table_borders(t_trib)
    doc.add_paragraph("")

    # 4. Itens Deduplicados (por descrição)
    if produtos_duplicados:
        doc.add_heading("4. Itens Deduplicados (por descrição)", level=1)
        t_dup = doc.add_table(rows=1, cols=4)
        t_dup.style = "Table Grid"
        hdr = t_dup.rows[0].cells
        hdr[0].text = "Código"
        hdr[1].text = "Descrição"
        hdr[2].text = "Ocorrências"
        hdr[3].text = "Valor Total"
        for c in hdr:
            c.paragraphs[0].runs[0].bold = True

        for item in produtos_duplicados:
            r = t_dup.add_row().cells
            r[0].text = str(item.get("codigo") or "")
            r[1].text = str(item.get("descricao") or "")
            r[2Claro ✅ — aqui estão os **arquivos completos e prontos para substituir no seu projeto**, com todas as correções implementadas:  

- 💯 Conversão correta de **alíquota decimal** (`10,2` → `0.102`)  
- 💸 Conversão automática de **imposto em centavos** (`336599` → `3365,99`)  
- 🧾 Deduplicação por **Código + Descrição**  
- 📊 Formatação correta de percentuais no DOCX  

---

## 📄 `app/services/analysis.py` (COMPLETO)

👉 [Clique para expandir/copiar todo o conteúdo]👇
````python
from __future__ import annotations
from datetime import datetime
from typing import Dict, Any
import io
import zipfile
import logging

from .nfe import parse_nfe_xml
from .ai_matcher import matcher

logger = logging.getLogger(__name__)

# -------------------------------------------------
# 🛡️ Funções auxiliares
# -------------------------------------------------
def _has_valid_ncm(ncm: str) -> bool:
    n = (ncm or '').strip()
    return len(n) == 8 and n.isdigit()

def parse_money_brl(value) -> float:
    """
    Converte strings ou números em float (R$).
    Aceita '3.365,99', '3365,99', '3365.99', 3365.99 ou 336599 (centavos).
    """
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        v = float(value)
    else:
        s = str(value).strip().replace("R$", "").replace(" ", "")
        if "," in s and s.count(",") == 1:
            s = s.replace(".", "")
            s = s.replace(",", ".")
        try:
            v = float(s)
        except Exception:
            return 0.0
    return float(v)

def parse_percent(value) -> float:
    """
    Converte entradas como '10,2', '10.2', '10', 10.2, 0.102 em fração (0.102).
    """
    if value is None:
        return None
    if isinstance(value, (int, float)):
        v = float(value)
    else:
        s = str(value).strip().replace("%", "").replace(" ", "").replace("R$", "")
        if "," in s and s.count(",") == 1:
            s = s.replace(".", "").replace(",", ".")
        try:
            v = float(s)
        except Exception:
            return None
    return v / 100.0 if v >= 1.0 else v

def safe_float(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0

def init_totals() -> dict:
    return {
        'documents': 0,
        'total_value_sum': 0.0,
        'period_start': None,
        'period_end': None,
        'items': 0,
        'revenue_excluded': 0.0,
        'revenue_excluded_breakdown': {},
        'st_correct_items_value': 0.0,
        'monofasico_palavra_chave': 0,
        'monofasico_sem_ncm': 0,
        'st_cfop_csosn_corretos': 0,
        'st_incorreta': 0,
        'erros_ncm_categoria': 0,
        'erros_outros': 0,
        'monofasico_total': 0,
        'monofasico_sem_cfop_csosn': 0,
        'tax_summary': {}
    }

# -------------------------------------------------
# 🧠 Função principal
# -------------------------------------------------
def run_analysis_from_bytes(zip_bytes: bytes, aliquota: float = None, imposto_pago: float = None) -> Dict[str, Any]:
    totals = init_totals()

    # 🔧 Normaliza entradas do usuário
    aliquota_frac = parse_percent(aliquota) if aliquota is not None else None
    imposto_pago_input = parse_money_brl(imposto_pago) if imposto_pago is not None else None

    produtos_raw = []
    dedup_map = {}
    excluidos = []

    with zipfile.ZipFile(io.BytesIO(zip_bytes), 'r') as zf:
        for name in zf.namelist():
            if not name.lower().endswith('.xml'):
                continue

            xml_bytes = zf.read(name)
            doc = parse_nfe_xml(xml_bytes)
            totals['documents'] += 1
            totals['total_value_sum'] += doc.get('total_value', 0.0)

            dt = doc.get('issue_date')
            if dt:
                if totals['period_start'] is None or dt < totals['period_start']:
                    totals['period_start'] = dt
                if totals['period_end'] is None or dt > totals['period_end']:
                    totals['period_end'] = dt

            for item in doc.get('items', []):
                totals['items'] += 1
                desc  = (item.get('xProd') or '').strip()
                ncm   = (item.get('ncm') or '').strip()
                cfop  = (item.get('cfop') or '').strip()
                csosn = (item.get('csosn') or '').strip()
                cprod = (item.get('cProd') or '').strip()
                vprod = float(item.get('vProd') or 0)

                # 👀 IA: Detectar monofásico
                is_mono = False
                hit = matcher.classify(desc)
                if hit and matcher.is_monofasico(hit[0]):
                    is_mono = True
                    totals['monofasico_palavra_chave'] += 1
                    totals['monofasico_total'] += 1
                    if not _has_valid_ncm(ncm):
                        totals['monofasico_sem_ncm'] += 1
                    if not (cfop and csosn):
                        totals['monofasico_sem_cfop_csosn'] += 1

                # ✅ ST Correto
                if is_mono and cfop == "5405" and csosn == "500":
                    totals['st_cfop_csosn_corretos'] += 1
                    totals['st_correct_items_value'] += vprod
                # ❌ ST Incorreto
                elif is_mono:
                    key = hit[0] if hit else "unknown"
                    totals['revenue_excluded_breakdown'].setdefault(key, 0.0)
                    totals['revenue_excluded_breakdown'][key] += vprod
                    totals['revenue_excluded'] += vprod
                    totals['st_incorreta'] += 1

                # 🚨 NCM vs categoria
                if hit:
                    val = matcher.validate_ncm_for_category(ncm, hit[0])
                    if not val["ncm_valido"]:
                        totals['erros_ncm_categoria'] += 1

                # 🧾 Guarda produto — apenas monofásico
                if is_mono:
                    prod_row = {
                        "descricao": desc,
                        "codigo": cprod,
                        "ncm": ncm,
                        "cfop": cfop,
                        "csosn": csosn,
                        "valor_total": vprod,
                        "valor_unitario": item.get('vUnCom'),
                        "quantidade": item.get('qCom'),
                        "numero": doc.get('numero'),
                        "data_emissao": dt.isoformat() if dt else None,
                        "chave": doc.get('chave'),
                        "monofasico": True,
                        "st_correto": (cfop == "5405" and csosn == "500"),
                        "ncm_recomendado": None,
                        "cest": item.get('cest') or None,
                        "cest_recomendado": None,
                    }
                    produtos_raw.append(prod_row)

                    # 📊 Deduplicação por código + descrição
                    key = (cprod.lower(), desc.lower())
                    if key not in dedup_map:
                        dedup_map[key] = {
                            "codigo": cprod,
                            "descricao": desc,
                            "ocorrencias": 1,
                            "valor_total": vprod,
                        }
                    else:
                        dedup_map[key]["ocorrencias"] += 1
                        dedup_map[key]["valor_total"] += vprod

                    if not prod_row["st_correto"]:
                        excluidos.append(prod_row)

    if totals['period_start']:
        totals['period_start'] = totals['period_start'].isoformat()
    if totals['period_end']:
        totals['period_end'] = totals['period_end'].isoformat()

    faturamento = totals['total_value_sum']
    receita_excluida = totals['revenue_excluded']
    base_corrigida = faturamento - receita_excluida

    imposto_corrigido = 0.0
    economia_estimada = 0.0
    imposto_pago_final = 0.0

    try:
        if aliquota_frac is not None:
            imposto_base_atual = faturamento * aliquota_frac
            imposto_corrigido = base_corrigida * aliquota_frac
            economia_estimada = max(0, imposto_base_atual - imposto_corrigido)
            imposto_pago_final = imposto_base_atual
        elif imposto_pago_input is not None:
            imposto_bruto = imposto_pago_input
            if faturamento > 0 and imposto_bruto > (faturamento * 3):
                imposto_bruto = imposto_bruto / 100.0

            aliquota_media = (imposto_bruto / faturamento) if faturamento > 0 else 0.0
            imposto_corrigido = base_corrigida * aliquota_media
            economia_estimada = max(0, imposto_bruto - imposto_corrigido)
            imposto_pago_final = imposto_bruto
    except Exception as e:
        logger.warning(f"[ANÁLISE] Falha no cálculo tributário: {e}")
        economia_estimada = 0.0
        imposto_corrigido = 0.0

    tax_summary = {
        'faturamento': faturamento,
        'base_corrigida': base_corrigida,
        'receita_excluida': receita_excluida,
        'imposto_corrigido': imposto_corrigido,
        'economia_estimada': economia_estimada,
        'aliquota_utilizada': (
            aliquota_frac if aliquota_frac is not None
            else ((imposto_pago_final / faturamento) if faturamento > 0 else 0.0)
        ),
        'imposto_pago': imposto_pago_final,
        'imposto_pago_informado': imposto_pago_input
    }

    for k, v in tax_summary.items():
        tax_summary[k] = safe_float(v)

    totals['tax_summary'] = tax_summary
    totals['products'] = produtos_raw
    totals['produtos_duplicados'] = sorted(dedup_map.values(), key=lambda x: x["valor_total"], reverse=True)
    totals['produtos_excluidos'] = excluidos

    logger.info(f"[DEBUG ANALYSIS] tax_summary final: {tax_summary}")
    logger.info(f"[DEBUG ANALYSIS] Monofásicos totais: {totals['monofasico_total']} / ST incorretos: {totals['st_incorreta']} / sem CFOP/CSOSN: {totals['monofasico_sem_cfop_csosn']}")
    return totals
